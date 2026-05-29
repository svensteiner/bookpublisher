from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Any


class ManuscriptReadError(RuntimeError):
    """A manuscript or supporting file could not be read.

    Carries a user-friendly German message that the GUI and CLI can
    show directly to a non-technical author - no Python traceback.
    """

    def __init__(self, path: Path | str, reason: str, hint: str) -> None:
        self.path = Path(path)
        self.reason = reason
        self.hint = hint
        message = (
            f"Datei konnte nicht gelesen werden:\n"
            f"  {self.path}\n\n"
            f"Grund: {reason}\n"
            f"So loest du das: {hint}"
        )
        super().__init__(message)


def _open_docx_document(path: Path) -> Any:
    """Open a DOCX with friendly errors. Raises ManuscriptReadError on failure."""

    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as exc:
        raise ManuscriptReadError(
            path,
            reason="Die Bibliothek python-docx ist nicht installiert.",
            hint=(
                "Oeffne ein Terminal im Programmordner und fuehre aus: "
                "pip install -r requirements.txt"
            ),
        ) from exc

    path_obj = Path(path)
    if not path_obj.exists():
        raise ManuscriptReadError(
            path_obj,
            reason="Die Datei wurde an diesem Pfad nicht gefunden.",
            hint=(
                "Pruefe, ob die Datei verschoben oder umbenannt wurde, "
                "und waehle den korrekten Buchordner erneut aus."
            ),
        )
    if path_obj.is_dir():
        raise ManuscriptReadError(
            path_obj,
            reason="Der Pfad zeigt auf einen Ordner statt auf eine DOCX-Datei.",
            hint="Waehle stattdessen die konkrete .docx-Datei innerhalb des Ordners aus.",
        )

    try:
        return Document(path_obj)
    except PackageNotFoundError as exc:
        raise ManuscriptReadError(
            path_obj,
            reason=(
                "Die Datei ist keine gueltige Word-DOCX-Datei "
                "(moeglicherweise beschaedigt oder im falschen Format)."
            ),
            hint=(
                "Oeffne die Datei in Word, speichere sie als .docx neu "
                "und starte die Pruefrunde danach erneut."
            ),
        ) from exc
    except zipfile.BadZipFile as exc:
        raise ManuscriptReadError(
            path_obj,
            reason="Die DOCX-Datei ist beschaedigt (defektes ZIP-Archiv).",
            hint=(
                "Stelle eine intakte Sicherung der Datei wieder her oder "
                "exportiere das Manuskript aus Word erneut als .docx."
            ),
        ) from exc
    except PermissionError as exc:
        raise ManuscriptReadError(
            path_obj,
            reason="Die Datei ist aktuell durch ein anderes Programm gesperrt.",
            hint=(
                "Schliesse Word oder andere Programme, die das Manuskript "
                "geoeffnet haben, und starte die Pruefrunde erneut."
            ),
        ) from exc
    except OSError as exc:
        raise ManuscriptReadError(
            path_obj,
            reason=f"Datei-System-Fehler beim Lesen der DOCX: {exc}",
            hint="Pruefe die Datei in Word und exportiere sie ggf. neu als .docx.",
        ) from exc


def read_docx_text(path: Path) -> str:
    doc = _open_docx_document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_pdf_text(path: Path, max_pages: int = 25) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise ManuscriptReadError(
            path,
            reason="Die Bibliothek PyMuPDF ist nicht installiert.",
            hint="Fuehre aus: pip install -r requirements.txt",
        ) from exc

    path_obj = Path(path)
    if not path_obj.exists():
        raise ManuscriptReadError(
            path_obj,
            reason="Die PDF-Datei wurde an diesem Pfad nicht gefunden.",
            hint="Pruefe den Buchordner und waehle ihn erneut aus.",
        )
    if path_obj.is_dir():
        raise ManuscriptReadError(
            path_obj,
            reason="Der Pfad zeigt auf einen Ordner statt auf eine PDF-Datei.",
            hint="Waehle stattdessen die konkrete .pdf-Datei innerhalb des Ordners aus.",
        )

    try:
        doc = fitz.open(path_obj)
    except ManuscriptReadError:
        raise
    except Exception as exc:
        raise ManuscriptReadError(
            path_obj,
            reason="Die PDF-Datei konnte nicht geoeffnet werden (vermutlich beschaedigt).",
            hint="Exportiere die PDF neu oder pruefe sie in einem PDF-Viewer.",
        ) from exc

    try:
        texts: list[str] = []
        for idx, page in enumerate(doc):
            if idx >= max_pages:
                break
            texts.append(page.get_text("text"))
        return "\n".join(texts)
    except ManuscriptReadError:
        raise
    except Exception as exc:
        raise ManuscriptReadError(
            path_obj,
            reason=f"Die Seiten der PDF-Datei konnten nicht gelesen werden: {exc}",
            hint="Exportiere die PDF neu oder pruefe sie in einem PDF-Viewer.",
        ) from exc
    finally:
        doc.close()


def read_text_file(path: Path) -> str:
    path_obj = Path(path)
    try:
        return path_obj.read_text(encoding="utf-8", errors="replace")
    except FileNotFoundError as exc:
        raise ManuscriptReadError(
            path_obj,
            reason="Die Textdatei wurde an diesem Pfad nicht gefunden.",
            hint="Pruefe den Buchordner und waehle ihn erneut aus.",
        ) from exc
    except PermissionError as exc:
        raise ManuscriptReadError(
            path_obj,
            reason="Die Textdatei ist aktuell durch ein anderes Programm gesperrt.",
            hint="Schliesse das blockierende Programm und starte die Pruefrunde erneut.",
        ) from exc
    except OSError as exc:
        raise ManuscriptReadError(
            path_obj,
            reason=f"Datei-System-Fehler beim Lesen der Datei: {exc}",
            hint="Pruefe Dateirechte und freien Speicherplatz auf dem Laufwerk.",
        ) from exc


def describe_zip(path: Path) -> dict:
    try:
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
        return {"path": str(path), "entries": names[:200], "entry_count": len(names)}
    except zipfile.BadZipFile:
        return {"path": str(path), "error": "bad_zip_file"}


def read_docx_chapters(path: Path) -> list:
    """Re-export of modules.chapters.extract_docx_chapters for convenience."""

    from modules.chapters import extract_docx_chapters

    return extract_docx_chapters(path)


def read_any_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return read_docx_text(path)
    if ext == ".pdf":
        return read_pdf_text(path)
    if ext in {".md", ".txt"}:
        return read_text_file(path)
    return ""


def open_docx_paragraphs(path: Path) -> Any:
    """Return an opened python-docx Document, raising ManuscriptReadError on failure.

    Shared entry point for modules that need access to the full paragraph or
    table iterator (chapters, industrial QA, sample scan) - keeps the
    user-friendly error handling in exactly one place.
    """

    return _open_docx_document(path)
