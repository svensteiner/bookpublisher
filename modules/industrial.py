from __future__ import annotations

import json
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from modules.cover import analyze_cover
from modules.discovery import BookProject
from modules.readers import read_any_text, read_text_file


@dataclass
class Gate:
    name: str
    status: str
    score: int
    findings: list[str] = field(default_factory=list)
    fixes: list[str] = field(default_factory=list)

    def to_json(self) -> dict[str, Any]:
        return {
            "name": self.name,
            "status": self.status,
            "score": self.score,
            "findings": self.findings,
            "fixes": self.fixes,
        }


def _status(score: int, blocking: bool = False) -> str:
    if blocking:
        return "FIX"
    if score >= 85:
        return "READY"
    if score >= 65:
        return "REVIEW"
    return "FIX"


def _word_count(text: str) -> int:
    return len(re.findall(r"\b[\wÄÖÜäöüß-]+\b", text, flags=re.UNICODE))


def _sentences(text: str) -> list[str]:
    return [item.strip() for item in re.split(r"(?<=[.!?])\s+", text) if item.strip()]


def _read_notes(project: BookProject) -> str:
    parts: list[str] = []
    for path in project.metadata_files + project.notes_files:
        if path.exists() and path.suffix.lower() in {".md", ".txt"}:
            try:
                parts.append(read_text_file(path))
            except OSError:
                continue
    return "\n\n".join(parts)


def _proof_pdfs(project: BookProject) -> list[Path]:
    if not project.root.exists():
        return []
    candidates = [path for path in project.root.rglob("*.pdf") if path.is_file()]
    return sorted(candidates, key=lambda path: str(path).lower())


def _extract_keywords(text: str) -> list[str]:
    patterns = [
        r"##\s*(?:Keywords|Suchbegriffe|Amazon Keywords|KDP Keywords)\s*(.*?)(?:\n##\s+|\Z)",
        r"(?:Keywords|Suchbegriffe)\s*:\s*(.+)",
    ]
    found: list[str] = []
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.I | re.S)
        if not match:
            continue
        block = match.group(1)
        for item in re.split(r"[\n,;]+", block):
            cleaned = item.strip(" -*\t")
            if 2 <= len(cleaned) <= 80:
                found.append(cleaned)
    return list(dict.fromkeys(found))[:12]


def analyze_docx_structure(path: Path | None) -> dict[str, Any]:
    if not path:
        return {"available": False}
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for industrial QA") from exc

    doc = Document(path)
    paragraphs: list[dict[str, Any]] = []
    headings: list[str] = []
    body_word_counts: list[int] = []
    non_empty_index = 0
    toc_detected = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        non_empty_index += 1
        style = para.style.name if para.style else ""
        words = _word_count(text)
        is_heading = "heading" in style.lower() or "überschrift" in style.lower()
        if is_heading:
            headings.append(text)
        else:
            body_word_counts.append(words)
        if non_empty_index <= 30 and re.search(r"\b(inhaltsverzeichnis|table of contents|toc)\b", text, flags=re.I):
            toc_detected = True
        paragraphs.append({"text": text, "style": style, "words": words, "is_heading": is_heading})

    text = "\n".join(item["text"] for item in paragraphs)
    total_words = _word_count(text)
    long_paragraphs = [item for item in paragraphs if not item["is_heading"] and item["words"] > 120]
    very_long_paragraphs = [item for item in paragraphs if not item["is_heading"] and item["words"] > 180]
    sample_words = max(1, round(total_words * 0.10))
    sample_text_parts: list[str] = []
    current_words = 0
    for item in paragraphs:
        sample_text_parts.append(item["text"])
        current_words += item["words"]
        if current_words >= sample_words:
            break
    sample_text = "\n".join(sample_text_parts)

    style_counts = Counter(item["style"] for item in paragraphs)
    return {
        "available": True,
        "path": str(path),
        "word_count": total_words,
        "paragraph_count": len(paragraphs),
        "body_paragraph_count": len(body_word_counts),
        "heading_count": len(headings),
        "first_headings": headings[:20],
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "toc_detected": toc_detected,
        "average_body_paragraph_words": round(sum(body_word_counts) / len(body_word_counts), 1) if body_word_counts else 0,
        "long_paragraph_count": len(long_paragraphs),
        "very_long_paragraph_count": len(very_long_paragraphs),
        "long_paragraph_ratio": round(len(long_paragraphs) / max(1, len(body_word_counts)), 3),
        "sample_word_target": sample_words,
        "sample_actual_words": current_words,
        "sample_sentence_count": len(_sentences(sample_text)),
        "sample_heading_count": sum(1 for item in paragraphs[: max(1, len(sample_text_parts))] if item["is_heading"]),
        "styles": dict(style_counts.most_common(20)),
    }


def _asset_gate(project: BookProject) -> Gate:
    findings: list[str] = []
    fixes: list[str] = []
    score = 100
    blocking = False

    required = {
        "manuscript_docx": project.manuscript,
        "cover_image": project.cover,
        "title": project.title,
        "subtitle": project.subtitle,
        "author": project.author,
        "amazon_description": project.amazon_description,
    }
    for name, value in required.items():
        if value:
            findings.append(f"{name}: present")
        else:
            findings.append(f"{name}: missing")
            fixes.append(f"Add production-ready {name}.")
            score -= 18
            if name in {"manuscript_docx", "cover_image", "title", "amazon_description"}:
                blocking = True

    return Gate("asset_completeness", _status(max(0, score), blocking), max(0, score), findings, fixes)


def _metadata_gate(project: BookProject, notes_text: str) -> Gate:
    findings: list[str] = []
    fixes: list[str] = []
    score = 100

    title_len = len(project.title or "")
    subtitle_len = len(project.subtitle or "")
    description = project.amazon_description or ""
    description_words = _word_count(description)
    keywords = _extract_keywords(notes_text)

    if 12 <= title_len <= 80:
        findings.append(f"title length: {title_len} chars")
    else:
        score -= 12
        fixes.append("Tune title length for Amazon scan speed and thumbnail comprehension.")

    if 20 <= subtitle_len <= 200:
        findings.append(f"subtitle length: {subtitle_len} chars")
    else:
        score -= 10
        fixes.append("Use the subtitle to state reader, method, and practical payoff.")

    if 120 <= description_words <= 650:
        findings.append(f"Amazon description: {description_words} words")
    else:
        score -= 16
        fixes.append("Bring Amazon description into a strong 120-650 word sales range.")

    if description.count("- ") >= 4 or description.count("\n-") >= 4:
        findings.append("description has scannable bullet structure")
    else:
        score -= 8
        fixes.append("Add buyer-friendly bullets to the Amazon description.")

    if len(keywords) >= 7:
        findings.append(f"keyword candidates found: {len(keywords)}")
    else:
        score -= 12
        fixes.append("Document at least 7 Amazon keyword candidates.")

    category_markers = re.findall(r"\b(kategorie|category|bisac|thema)\b", notes_text, flags=re.I)
    if category_markers:
        findings.append("category notes found")
    else:
        score -= 10
        fixes.append("Document primary and secondary KDP category hypotheses.")

    return Gate("metadata_and_storefront", _status(max(0, score)), max(0, score), findings, fixes)


def _kindle_gate(profile: dict[str, Any]) -> Gate:
    findings: list[str] = []
    fixes: list[str] = []
    score = 100
    blocking = False

    if not profile.get("available"):
        return Gate("kindle_ebook_readiness", "FIX", 0, ["manuscript not available"], ["Add a DOCX manuscript."])

    words = int(profile.get("word_count") or 0)
    headings = int(profile.get("heading_count") or 0)
    avg_para = float(profile.get("average_body_paragraph_words") or 0)
    long_ratio = float(profile.get("long_paragraph_ratio") or 0)

    findings.append(f"word count: {words}")
    findings.append(f"heading count: {headings}")
    findings.append(f"average body paragraph words: {avg_para}")
    findings.append(f"long paragraph ratio: {long_ratio}")

    if words < 12000:
        score -= 18
        fixes.append("Confirm the ebook has enough perceived value for paid nonfiction.")
    if headings < 8:
        score -= 18
        fixes.append("Add/verify a clear heading hierarchy for Kindle navigation.")
    if avg_para > 75:
        score -= 12
        fixes.append("Shorten average paragraph length for mobile Kindle reading.")
    if long_ratio > 0.12:
        score -= 12
        fixes.append("Break long paragraphs before Kindle upload.")
    if profile.get("table_count", 0) > 0:
        score -= 5
        fixes.append("Preview every table on phone-sized Kindle layouts.")
    if not profile.get("toc_detected"):
        score -= 8
        fixes.append("Verify the DOCX heading styles generate a clickable Kindle TOC.")
    if profile.get("sample_sentence_count", 0) < 20:
        score -= 8
        fixes.append("Strengthen the first 10 percent sample with clearer promise and proof.")

    if words == 0:
        blocking = True
    return Gate("kindle_ebook_readiness", _status(max(0, score), blocking), max(0, score), findings, fixes)


def _production_gate(project: BookProject, profile: dict[str, Any]) -> Gate:
    findings: list[str] = []
    fixes: list[str] = []
    score = 100
    blocking = False

    if project.cover:
        cover = analyze_cover(project.cover)
        findings.append(f"cover size: {cover['width']} x {cover['height']} px")
        findings.append(f"cover ratio: {cover['aspect_ratio_height_width']}")
        if cover["width"] < 1500 or cover["height"] < 2500:
            score -= 20
            fixes.append("Export ebook cover at production size before KDP upload.")
        if not 1.45 <= cover["aspect_ratio_height_width"] <= 1.65:
            score -= 20
            fixes.append("Adjust ebook cover ratio near 1.6 height/width.")
        if cover["low_thumbnail_contrast_risk"]:
            score -= 12
            fixes.append("Increase thumbnail contrast for Amazon search results.")
        if cover["very_light_cover"]:
            score -= 8
            fixes.append("Check edge visibility on Amazon white backgrounds.")
    else:
        score -= 35
        blocking = True
        fixes.append("Add a production-ready cover image.")

    if profile.get("inline_shape_count", 0) > 0:
        findings.append(f"inline images in manuscript: {profile['inline_shape_count']}")
        fixes.append("Preview every inline image in Kindle Previewer.")
        score -= 5
    else:
        findings.append("inline images in manuscript: 0")

    proof_pdfs = project.pdf_files or _proof_pdfs(project)
    if proof_pdfs:
        findings.append(f"PDF/proof exports found: {len(proof_pdfs)}")
    else:
        score -= 5
        fixes.append("Keep a rendered customer PDF/proof outside upload-critical files for final visual QA.")

    return Gate("production_package", _status(max(0, score), blocking), max(0, score), findings, fixes)


def _sellability_gate(project: BookProject, profile: dict[str, Any], notes_text: str) -> Gate:
    findings: list[str] = []
    fixes: list[str] = []
    score = 100
    combined = "\n".join([project.title or "", project.subtitle or "", project.amazon_description or "", notes_text])

    markers = {
        "specific_reader": r"\b(selbstst[aä]ndige|gr[uü]nder|manager|cfo|berater|wissensarbeiter)\b",
        "anti_hype_positioning": r"\b(kein hype|keine tool-liste|keine motivations|n[üu]chtern|feldnotiz|echte fehler)\b",
        "proof": r"\b(unter 50|90 tage|47 linkedin|keine zahlenden kunden|tradingbot|prompt-leak)\b",
        "practical_payoff": r"\b(startpunkte|kontrollregeln|aufgaben|praktisch|abgeben|bleiben)\b",
    }
    for name, pattern in markers.items():
        if re.search(pattern, combined, flags=re.I):
            findings.append(f"{name}: present")
        else:
            score -= 12
            fixes.append(f"Make {name.replace('_', ' ')} explicit on the product page.")

    if profile.get("sample_heading_count", 0) >= 2:
        findings.append("first 10 percent has structural signposts")
    else:
        score -= 10
        fixes.append("Make the first Kindle sample structurally obvious within the first 10 percent.")

    if profile.get("sample_sentence_count", 0) >= 35:
        findings.append("first 10 percent has enough reading surface")
    else:
        score -= 8
        fixes.append("Ensure the first 10 percent contains enough concrete promise, proof, and payoff.")

    return Gate("amazon_sellability", _status(max(0, score)), max(0, score), findings, fixes)


def build_industrial_qa(project: BookProject, agent_context: dict[str, Any] | None = None) -> dict[str, Any]:
    notes_text = _read_notes(project)
    profile = analyze_docx_structure(project.manuscript)
    gates = [
        _asset_gate(project),
        _metadata_gate(project, notes_text),
        _kindle_gate(profile),
        _production_gate(project, profile),
        _sellability_gate(project, profile, notes_text),
    ]
    blocking = [gate for gate in gates if gate.status == "FIX"]
    review = [gate for gate in gates if gate.status == "REVIEW" or gate.fixes]
    average_score = round(sum(gate.score for gate in gates) / len(gates)) if gates else 0
    decision = "GO" if not blocking and not review and average_score >= 85 else "GO_AFTER_FIXES" if not blocking else "HOLD"
    investor_grade = round(average_score / 10, 1)

    return {
        "project_id": project.project_id,
        "decision": decision,
        "industrial_score": average_score,
        "investor_grade": investor_grade,
        "gates": [gate.to_json() for gate in gates],
        "docx_profile": profile,
        "keywords_found": _extract_keywords(notes_text),
        "all_required_fixes": [fix for gate in gates for fix in gate.fixes],
        "agent_context": agent_context or {},
    }


def _gate_lookup(report: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {gate["name"]: gate for gate in report.get("gates", [])}


def _traffic_light(decision: str) -> tuple[str, str]:
    if decision == "GO":
        return "GRUEN", "Du kannst nach einer normalen Sichtpruefung hochladen."
    if decision == "GO_AFTER_FIXES":
        return "GELB", "Das Buch ist stark, aber vor dem Upload bleiben einzelne Kontrollpunkte."
    return "ROT", "Nicht hochladen. Es gibt blockierende Punkte."


def _affected_file_for_fix(fix: str) -> str:
    lower = fix.lower()
    if "cover" in lower or "thumbnail" in lower or "background" in lower:
        return "Cover-Datei"
    if "description" in lower or "keyword" in lower or "category" in lower or "title" in lower or "subtitle" in lower:
        return "Amazon_Beschreibung_und_Metadaten.md"
    if "docx" in lower or "kindle" in lower or "paragraph" in lower or "toc" in lower or "sample" in lower or "ebook" in lower:
        return "DOCX-Manuskript"
    if "pdf" in lower or "proof" in lower:
        return "Proof-PDF"
    return "Buchordner"


def _plain_fix(fix: str) -> str:
    lower = fix.lower()
    if "enough perceived value" in lower:
        return "Pruefe, ob das Buch fuer den geplanten Preis genug wahrgenommenen Wert bietet. Bei rund 9.000 Woertern sollte der Nutzen sehr klar sein."
    if "edge visibility" in lower:
        return "Pruefe, ob das Cover auf weissem Amazon-Hintergrund noch klar abgegrenzt ist."
    if "kindle previewer" in lower:
        return "Oeffne die Datei in der Kindle-Vorschau und pruefe Handy-, Tablet- und Kindle-Ansicht."
    if "clickable kindle toc" in lower or "toc" in lower:
        return "Pruefe, ob das Inhaltsverzeichnis in Kindle klickbar ist."
    if "amazon description" in lower:
        return "Passe die Amazon-Beschreibung an, damit Nutzen und Zielgruppe sofort klar sind."
    if "keyword" in lower:
        return "Ergaenze mindestens sieben passende Amazon-Suchbegriffe."
    if "category" in lower:
        return "Lege passende Amazon-Kategorien fest."
    if "paragraph" in lower:
        return "Kuerze zu lange Absaetze, damit das Buch am Handy leichter lesbar ist."
    if "cover" in lower:
        return "Pruefe oder ersetze die Cover-Datei."
    return fix


def render_beginner_summary(project: BookProject, report: dict[str, Any]) -> str:
    decision = str(report.get("decision", "HOLD"))
    light, plain_decision = _traffic_light(decision)
    gates = _gate_lookup(report)
    profile = report.get("docx_profile", {})
    fixes = report.get("all_required_fixes", [])

    lines = [
        "# Einfache Buch-Pruefung",
        "",
        f"Buch: **{project.title or project.project_id}**",
        f"Ampel: **{light}**",
        f"Ergebnis: {plain_decision}",
        f"Score: **{report.get('industrial_score', 'n/a')}/100**",
        "",
        "## Was ist schon gut?",
        "",
    ]

    positives: list[str] = []
    if gates.get("asset_completeness", {}).get("status") == "READY":
        positives.append("Alle wichtigen Dateien sind da: Manuskript, Cover, Titel, Autor und Amazon-Beschreibung.")
    if gates.get("metadata_and_storefront", {}).get("status") == "READY":
        positives.append("Die Amazon-Seite ist grundsaetzlich vorbereitet: Beschreibung, Keywords und Kategorien sind vorhanden.")
    if gates.get("amazon_sellability", {}).get("status") == "READY":
        positives.append("Das Buch hat eine klare Zielgruppe, konkrete Beweise und eine verkaufbare Positionierung.")
    if gates.get("production_package", {}).get("score", 0) >= 85:
        positives.append("Cover-Format und Produktionspaket sehen technisch gut aus.")
    if profile.get("heading_count", 0) >= 8:
        positives.append("Das Manuskript hat genug Ueberschriften fuer Kindle-Navigation und Lesbarkeit.")

    lines.extend(f"- {item}" for item in positives or ["Die Grundstruktur wurde erkannt und kann geprueft werden."])

    lines.extend(["", "## Was musst du jetzt tun?", ""])
    if fixes:
        for idx, fix in enumerate(fixes, start=1):
            lines.extend([
                f"{idx}. {_plain_fix(fix)}",
                f"   Betroffene Datei: **{_affected_file_for_fix(fix)}**",
            ])
    else:
        lines.append("Nichts Blockierendes. Mache nur noch eine menschliche Endkontrolle.")

    lines.extend([
        "",
        "## Was bedeutet das praktisch?",
        "",
    ])
    if decision == "GO":
        lines.append("Du kannst den KDP-Upload vorbereiten. Pruefe trotzdem einmal die Kindle-Vorschau und die Amazon-Produktseite.")
    elif decision == "GO_AFTER_FIXES":
        lines.append("Du bist nah an der Veroeffentlichung. Arbeite die Punkte oben ab, speichere die Dateien, und starte danach die naechste Pruefrunde.")
    else:
        lines.append("Bitte zuerst die fehlenden oder blockierenden Dateien korrigieren. Danach eine neue Pruefrunde starten.")

    lines.extend([
        "",
        "## Erkannte Dateien",
        "",
        f"- Manuskript: `{project.manuscript or 'FEHLT'}`",
        f"- Cover: `{project.cover or 'FEHLT'}`",
        f"- Metadaten-Dateien: {len(project.metadata_files)}",
        "",
        "## Naechster Klick",
        "",
        "Nach deinen Anpassungen wieder **Pruefrunde starten** klicken.",
    ])
    return "\n".join(lines)


def render_industrial_qa_markdown(report: dict[str, Any]) -> str:
    lines = [
        "# Industrial Publisher QA",
        "",
        f"Project: `{report['project_id']}`",
        f"Decision: **{report['decision']}**",
        f"Industrial score: **{report['industrial_score']}/100**",
        f"Investor grade: **{report['investor_grade']}/10**",
        "",
    ]
    agent_context = report.get("agent_context") or {}
    skills = agent_context.get("skills") or []
    memory = agent_context.get("memory") or {}
    if skills or memory:
        lines.extend(["## Agent System", ""])
        if skills:
            lines.append("Loaded skills:")
            lines.extend(f"- {skill.get('name', 'unknown')}: {skill.get('purpose', 'n/a')}" for skill in skills)
            lines.append("")
        project_memory = memory.get("project_memory") or {}
        if project_memory.get("rounds"):
            latest = project_memory["rounds"][-1]
            lines.extend([
                "Memory:",
                f"- Previous decision: {latest.get('decision', 'n/a')}",
                f"- Previous industrial score: {latest.get('industrial_score', 'n/a')}",
                "",
            ])
    lines.extend([
        "## Gates",
        "",
    ])
    for gate in report["gates"]:
        lines.extend([
            f"### {gate['name']} - {gate['status']} ({gate['score']}/100)",
            "",
            "Findings:",
            *(f"- {item}" for item in gate["findings"]),
            "",
        ])
        if gate["fixes"]:
            lines.extend(["Required fixes:", *(f"- {item}" for item in gate["fixes"]), ""])

    profile = report["docx_profile"]
    if profile.get("available"):
        lines.extend([
            "## Kindle Profile",
            "",
            f"- Word count: {profile['word_count']}",
            f"- Paragraphs: {profile['paragraph_count']}",
            f"- Headings: {profile['heading_count']}",
            f"- Tables: {profile['table_count']}",
            f"- Inline images: {profile['inline_shape_count']}",
            f"- Average body paragraph words: {profile['average_body_paragraph_words']}",
            f"- Long paragraph ratio: {profile['long_paragraph_ratio']}",
            f"- TOC detected in manuscript text: {profile['toc_detected']}",
            "",
        ])
    if report["all_required_fixes"]:
        lines.extend([
            "## Release-Critical Fixes",
            "",
            *(f"{idx}. {item}" for idx, item in enumerate(report["all_required_fixes"], start=1)),
            "",
        ])
    lines.extend([
        "## Machine JSON",
        "",
        "```json",
        json.dumps(
            {
                "decision": report["decision"],
                "industrial_score": report["industrial_score"],
                "investor_grade": report["investor_grade"],
            },
            ensure_ascii=False,
            indent=2,
        ),
        "```",
    ])
    return "\n".join(lines)
